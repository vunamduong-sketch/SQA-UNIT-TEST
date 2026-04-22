from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_FILE = "agoda-be-model-test-matrix.docx"

rows = [
    # accounts
    ("agoda-be|accounts", "CustomUser", "agoda-be\\accounts\\models.py", "__str__", "Returns username; core identity shown across UI/logs."),

    # activities
    ("agoda-be|activities", "Activity", "agoda-be\\activities\\models.py", "save", "Recomputes weighted score on every save; affects ranking/recommendation output."),
    ("agoda-be|activities", "UserActivityInteraction", "agoda-be\\activities\\models.py", "update_weighted_score", "Updates personalized user-activity score for recommendation logic."),
    ("agoda-be|activities", "ActivityDate", "agoda-be\\activities\\models.py", "get_active_promotion", "Resolves active promotion and best discount for pricing."),
    ("agoda-be|activities", "ActivityDateBookingDetail", "agoda-be\\activities\\models.py", "save", "Calculates total/final price, updates participants, propagates booking totals."),

    # airlines
    ("agoda-be|airlines", "Airline", "agoda-be\\airlines\\models.py", "__str__", "Ensures airline display text is correct (name + code)."),
    ("agoda-be|airlines", "Aircraft", "agoda-be\\airlines\\models.py", "__str__", "Ensures aircraft display text includes airline code/model/registration."),

    # cars
    ("agoda-be|cars", "Car", "agoda-be\\cars\\models.py", "calc_total_weighted_score", "Core ranking formula for car recommendation ordering."),
    ("agoda-be|cars", "Car", "agoda-be\\cars\\models.py", "update_total_weighted_score", "Persists weighted score used in sorting/filtering."),
    ("agoda-be|cars", "Car", "agoda-be\\cars\\models.py", "get_active_promotion", "Finds valid best promotion before price calculation."),
    ("agoda-be|cars", "CarBookingDetail", "agoda-be\\cars\\models.py", "save", "Assigns driver, updates driver status, computes total/final price, syncs booking totals."),
    ("agoda-be|cars", "UserCarInteraction", "agoda-be\\cars\\models.py", "update_weighted_score", "Updates user-car affinity score for personalization."),

    # chatbots
    ("agoda-be|chatbots", "(No custom model)", "agoda-be\\chatbots\\models.py", "N/A", "Module currently has no model class/method to test."),

    # chats
    ("agoda-be|chats", "Conversation", "agoda-be\\chats\\models.py", "__str__", "Verifies readable conversation label for admin/debugging."),
    ("agoda-be|chats", "Message", "agoda-be\\chats\\models.py", "__str__", "Verifies message label includes sender and conversation id."),

    # handbooks
    ("agoda-be|handbooks", "Handbook", "agoda-be\\handbooks\\models.py", "save", "Recomputes weighted score; impacts handbook ranking."),
    ("agoda-be|handbooks", "Handbook", "agoda-be\\handbooks\\models.py", "update_total_weighted_score", "Persists aggregate score for recommendation/filtering."),
    ("agoda-be|handbooks", "UserHandbookInteraction", "agoda-be\\handbooks\\models.py", "update_weighted_score", "Updates user-handbook personalized relevance."),

    # hotels
    ("agoda-be|hotels", "Hotel", "agoda-be\\hotels\\models.py", "save", "Keeps weighted score consistent for hotel ranking."),
    ("agoda-be|hotels", "Hotel", "agoda-be\\hotels\\models.py", "update_min_price", "Calculates displayed minimum/average available room price."),
    ("agoda-be|hotels", "UserHotelInteraction", "agoda-be\\hotels\\models.py", "update_weighted_score", "Updates personalized hotel score from interactions."),

    # images
    ("agoda-be|images", "Image", "agoda-be\\images\\models.py", "__str__", "Ensures stored file name/path is rendered correctly."),

    # payments
    ("agoda-be|payments", "Payment", "agoda-be\\payments\\models.py", "(model default: status)", "Critical payment flow default state must start as Pending."),

    # promotions
    ("agoda-be|promotions", "Promotion", "agoda-be\\promotions\\models.py", "__str__", "Verifies promotion label with promotion type display."),
    ("agoda-be|promotions", "RoomPromotion", "agoda-be\\promotions\\models.py", "__str__", "Ensures room-promotion linkage text is correct."),
    ("agoda-be|promotions", "CarPromotion", "agoda-be\\promotions\\models.py", "__str__", "Ensures car-promotion linkage text handles nullable car and named car."),
    ("agoda-be|promotions", "ActivityPromotion", "agoda-be\\promotions\\models.py", "__str__", "Ensures activity date linkage is traceable for debugging."),

    # reviews
    ("agoda-be|reviews", "Review", "agoda-be\\reviews\\models.py", "service_type_name", "Returns service label used in response and display logic."),
    ("agoda-be|reviews", "Review", "agoda-be\\reviews\\models.py", "get_service_instance", "Maps review to concrete service object (hotel/activity)."),
    ("agoda-be|reviews", "Review", "agoda-be\\reviews\\models.py", "__str__", "Produces readable review identity containing service/ref id."),

    # rooms
    ("agoda-be|rooms", "Room", "agoda-be\\rooms\\models.py", "capacity", "Computes total room capacity from adults + children."),
    ("agoda-be|rooms", "Room", "agoda-be\\rooms\\models.py", "save", "Updates availability state and triggers hotel min price refresh."),
    ("agoda-be|rooms", "Room", "agoda-be\\rooms\\models.py", "decrease_available_rooms", "Prevents negative inventory and updates availability after booking."),
    ("agoda-be|rooms", "Room", "agoda-be\\rooms\\models.py", "get_active_promotion", "Finds current best room-level promotion for final pricing."),
    ("agoda-be|rooms", "RoomBookingDetail", "agoda-be\\rooms\\models.py", "save", "Calculates booking amount (overnight/dayuse), applies discount, decrements stock, syncs booking totals."),
]


def main():
    doc = Document()

    title = doc.add_heading("Files / Classes / Functions THAT ARE Tested", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("agoda-be (accounts, activities, airlines, cars, chatbots, chats, handbooks, hotels, images, payments, promotions, reviews, rooms)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "No."
    hdr[1].text = "File / Class"
    hdr[2].text = "Path"
    hdr[3].text = "Function / Method"
    hdr[4].text = "Reason"

    for idx, (group_name, cls_name, path, method, reason) in enumerate(rows, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = f"{group_name}\n{cls_name}"
        row[2].text = path
        row[3].text = method
        row[4].text = reason

    doc.save(OUTPUT_FILE)
    print(f"Generated: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
